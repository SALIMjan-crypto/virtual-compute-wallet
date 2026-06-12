# -*- coding: utf-8 -*-
"""
Virtual Compute Wallet — zypl.ai × Zinda Capital
(formerly GPU Credit Card)

v2.0 — Post-meeting update (Akai Akmal, June 2 2026)
Pivot: Physical co-branded card → Embedded GPU Credit Module in Zinda Internet Banking

Generator: Product Specifications, CJM, Business Model, Flow Diagram
"""

import os
import io
import math
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.patheffects as pe
import numpy as np
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Output directory ───
OUTPUT_DIR = r"c:\ANTIGRAVITY PROJECTS\GPU Credit Card"
CHARTS_DIR = os.path.join(OUTPUT_DIR, "charts")
os.makedirs(CHARTS_DIR, exist_ok=True)

# ─── Brand colors ───
BRAND_PRIMARY = "#1A1A2E"      # Deep navy
BRAND_ACCENT = "#E94560"       # Vibrant red-pink
BRAND_GREEN = "#00B894"        # Success green
BRAND_BLUE = "#0984E3"         # Tech blue
BRAND_GOLD = "#FDCB6E"         # Gold accent
BRAND_PURPLE = "#6C5CE7"       # Purple
BRAND_DARK = "#0F0F1A"         # Almost black
BRAND_LIGHT = "#F5F6FA"        # Light gray
BRAND_GRADIENT_1 = "#667EEA"   # Gradient start
BRAND_GRADIENT_2 = "#764BA2"   # Gradient end
BRAND_TEAL = "#00CEC9"         # Teal accent


def hex_to_rgb(hex_color):
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex.lstrip("#")}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_formatted_paragraph(doc, text, font_name='Calibri', size=11,
                             bold=False, italic=False, color=None,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=0, space_after=6):
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        r, g, b = hex_to_rgb(color)
        run.font.color.rgb = RGBColor(r, g, b)
    return p


def add_section_title(doc, text, level=1):
    """Add a styled section title."""
    colors_by_level = {
        1: BRAND_PRIMARY,
        2: BRAND_BLUE,
        3: BRAND_PURPLE,
    }
    sizes_by_level = {1: 22, 2: 16, 3: 13}
    color = colors_by_level.get(level, BRAND_PRIMARY)
    size = sizes_by_level.get(level, 14)
    
    p = add_formatted_paragraph(
        doc, text, font_name='Calibri', size=size,
        bold=True, color=color,
        space_before=18 if level == 1 else 12,
        space_after=8
    )
    
    # Add bottom border for level 1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="{BRAND_ACCENT.lstrip("#")}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    return p


def add_styled_table(doc, headers, rows, header_color=BRAND_PRIMARY,
                      accent_row_color="#F0F3FF", col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    # Header row
    hdr = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[1 + row_idx]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            if row_idx % 2 == 0:
                set_cell_shading(cell, accent_row_color)
    
    return table


# ═══════════════════════════════════════════════════════════════
#  CHART GENERATION FUNCTIONS
# ═══════════════════════════════════════════════════════════════

def create_revenue_projection_chart():
    """Revenue projection chart."""
    fig, ax = plt.subplots(figsize=(10, 5))
    fig.patch.set_facecolor('#FAFBFE')
    
    categories = ['4 GPU\n(Pilot)', '16 GPU\n(Phase 1)', '50 GPU\n(Phase 2)', '100 GPU\n(Scale)']
    min_rev = [90, 360, 1125, 2250]
    max_rev = [170, 680, 2125, 4250]
    
    x = np.arange(len(categories))
    width = 0.35
    
    bars1 = ax.bar(x - width/2, min_rev, width, label='Min. Revenue ($K/yr)',
                   color=BRAND_BLUE, alpha=0.85, edgecolor='white', linewidth=0.5)
    bars2 = ax.bar(x + width/2, max_rev, width, label='Max. Revenue ($K/yr)',
                   color=BRAND_ACCENT, alpha=0.85, edgecolor='white', linewidth=0.5)
    
    for bar in bars1:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height + 30,
                f'${int(height)}K', ha='center', va='bottom',
                fontsize=9, fontweight='bold', color=BRAND_BLUE)
    
    for bar in bars2:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height + 30,
                f'${int(height)}K', ha='center', va='bottom',
                fontsize=9, fontweight='bold', color=BRAND_ACCENT)
    
    ax.set_ylabel('Revenue (USD K / year)', fontsize=11, fontweight='bold', color='#333')
    ax.set_title('Revenue Projection by Scaling Phase',
                 fontsize=14, fontweight='bold', color=BRAND_PRIMARY, pad=15)
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=10)
    ax.legend(frameon=True, framealpha=0.9, fontsize=9)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_ylim(0, 5000)
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    
    plt.tight_layout()
    path = os.path.join(CHARTS_DIR, 'revenue_projection.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def create_gpu_utilization_gauge():
    """GPU utilization and capacity gauge chart."""
    fig, axes = plt.subplots(1, 3, figsize=(12, 4))
    fig.patch.set_facecolor('#FAFBFE')
    
    metrics = [
        ('GPU Pool\n(GPU-hrs/mo)', 2880, 'GPU-hrs', BRAND_BLUE),
        ('HBM3e Memory\n(GB per GPU)', 141, 'GB', BRAND_PURPLE),
        ('Max Sessions\n(concurrent)', 4, 'sessions', BRAND_GREEN),
    ]
    
    for ax, (label, value, unit, color) in zip(axes, metrics):
        ax.set_xlim(-1.5, 1.5)
        ax.set_ylim(-1.5, 1.5)
        ax.set_aspect('equal')
        ax.axis('off')
        
        # Outer ring
        theta = np.linspace(0, 2*np.pi, 100)
        ax.plot(np.cos(theta), np.sin(theta), color='#E0E0E0', linewidth=8, alpha=0.4)
        
        # Filled arc (~80%)
        fill_pct = 0.80
        theta_fill = np.linspace(np.pi/2, np.pi/2 - 2*np.pi*fill_pct, 100)
        ax.plot(np.cos(theta_fill), np.sin(theta_fill), color=color, linewidth=10, alpha=0.8)
        
        # Center text
        ax.text(0, 0.15, f'{value:,}', ha='center', va='center',
                fontsize=22, fontweight='bold', color=color)
        ax.text(0, -0.2, unit, ha='center', va='center',
                fontsize=10, color='#666')
        ax.text(0, -1.35, label, ha='center', va='center',
                fontsize=9, fontweight='bold', color='#333', linespacing=1.3)
    
    fig.suptitle('Key Infrastructure Metrics — NVIDIA H200',
                 fontsize=13, fontweight='bold', color=BRAND_PRIMARY, y=1.02)
    plt.tight_layout()
    path = os.path.join(CHARTS_DIR, 'gpu_metrics.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def create_business_model_canvas():
    """Business Model Canvas visual — updated for Virtual Compute Wallet."""
    fig, ax = plt.subplots(figsize=(14, 9))
    fig.patch.set_facecolor('#FAFBFE')
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')
    
    # Title
    ax.text(7, 8.6, 'BUSINESS MODEL CANVAS — VIRTUAL COMPUTE WALLET',
            ha='center', fontsize=15, fontweight='bold', color=BRAND_PRIMARY,
            path_effects=[pe.withStroke(linewidth=0, foreground='white')])
    
    # Canvas blocks: (x, y, w, h, title, content_lines, color)
    blocks = [
        # Top row
        (0.2, 4.5, 2.6, 3.8, 'KEY\nPARTNERS',
         ['- zypl.ai\n  (GPU infra +\n  scoring)', '- Zinda Capital\n  (IB platform,\n  wallet)',
          '- NVIDIA\n  (hardware)', '- Central Bank\n  of Tajikistan'],
         BRAND_PURPLE),
        
        (3.0, 6.3, 2.6, 2.0, 'KEY\nACTIVITIES',
         ['- GPU fleet mgmt', '- Metering/billing', '- Wallet operations', '- API integration'],
         BRAND_BLUE),
        
        (3.0, 4.5, 2.6, 1.6, 'KEY\nRESOURCES',
         ['- 4x NVIDIA H200', '- Zinda IB platform', '- zypl.ai scoring'],
         '#2196F3'),
        
        (5.8, 4.5, 2.6, 3.8, 'VALUE\nPROPOSITION',
         ['* In-app GPU\n   access',
          '* TJS wallet\n   (no FX friction)',
          '* Pay-per-use\n   GPU-hours',
          '* Sovereign\n   compute'],
         BRAND_ACCENT),
        
        (8.6, 6.3, 2.6, 2.0, 'CUSTOMER\nRELATIONSHIPS',
         ['- Self-service app', '- API access', '- 24/7 support'],
         BRAND_GREEN),
        
        (8.6, 4.5, 2.6, 1.6, 'CHANNELS',
         ['- Mobile Wallet App', '- Internet Banking', '- REST API'],
         '#26A69A'),
        
        (11.4, 4.5, 2.4, 3.8, 'CUSTOMER\nSEGMENTS',
         ['Mobile Wallet:', '- Researchers, Profs', '- Universities',
          'Internet Banking:', '- AI startups', '- Fintechs, Corps'],
         '#FF7043'),
        
        # Bottom row
        (0.2, 0.5, 6.6, 3.5, 'COST STRUCTURE',
         ['- H200 amortization ($120-160K per GPU)',
          '- Electricity & cooling',
          '- Zinda platform operations',
          '- Network infrastructure',
          '- zypl.ai platform development'],
         '#546E7A'),
        
        (7.0, 0.5, 6.8, 3.5, 'REVENUE STREAMS',
         ['$ GPU-hour billing: $3-4/h (H200)',
          '$ Pilot (4 GPU): $90K-$170K/yr',
          '$ Scale (100 GPU): $2.25M-$4.25M/yr',
          '$ Wallet service fees',
          '$ Premium SLA & priority access'],
         '#43A047'),
    ]
    
    for (x, y, w, h, title, lines, color) in blocks:
        # Draw rounded rectangle
        rect = FancyBboxPatch(
            (x, y), w, h,
            boxstyle="round,pad=0.1",
            facecolor='white', edgecolor=color,
            linewidth=2, alpha=0.95
        )
        ax.add_patch(rect)
        
        # Title bar
        title_rect = FancyBboxPatch(
            (x, y + h - 0.7), w, 0.7,
            boxstyle="round,pad=0.05",
            facecolor=color, edgecolor=color,
            linewidth=1, alpha=0.9
        )
        ax.add_patch(title_rect)
        
        ax.text(x + w/2, y + h - 0.35, title,
                ha='center', va='center', fontsize=7.5,
                fontweight='bold', color='white', linespacing=1.1)
        
        # Content
        content_y = y + h - 1.0
        for line in lines:
            ax.text(x + 0.15, content_y, line,
                    ha='left', va='top', fontsize=6.5,
                    color='#333', linespacing=1.2)
            content_y -= 0.55
    
    plt.tight_layout()
    path = os.path.join(CHARTS_DIR, 'business_model_canvas.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def create_cjm_diagram():
    """Customer Journey Map — updated for embedded module (no physical card)."""
    fig, ax = plt.subplots(figsize=(14, 8))
    fig.patch.set_facecolor('#FAFBFE')
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    
    ax.text(7, 7.6, 'CUSTOMER JOURNEY MAP (CJM) — VIRTUAL COMPUTE WALLET',
            ha='center', fontsize=14, fontweight='bold', color=BRAND_PRIMARY)
    ax.text(7, 7.2, 'Embedded Module in Zinda Internet Banking',
            ha='center', fontsize=10, fontstyle='italic', color='#777')
    
    # Journey stages — updated for in-app flow
    stages = [
        {
            'name': '1. DISCOVER',
            'emoji': '>>',
            'touchpoints': 'Zinda App banner,\nNotifications, Web',
            'actions': 'Sees "Cloud Compute"\nmodule in Zinda app',
            'emotions': 'Curiosity',
            'pain': 'Unaware of\nlocal GPU option',
            'color': BRAND_BLUE
        },
        {
            'name': '2. ACTIVATE',
            'emoji': '>>',
            'touchpoints': 'In-app module,\nTerms & Conditions',
            'actions': 'Opens module,\naccepts T&C',
            'emotions': 'Interest',
            'pain': 'Understanding\nGPU pricing',
            'color': BRAND_PURPLE
        },
        {
            'name': '3. FUND',
            'emoji': '>>',
            'touchpoints': 'Wallet top-up,\nCredit application',
            'actions': 'Tops up TJS wallet\nor applies for credit',
            'emotions': 'Anticipation',
            'pain': 'TJS to GPU-hr\nconversion clarity',
            'color': '#FF6F00'
        },
        {
            'name': '4. LAUNCH',
            'emoji': '>>',
            'touchpoints': '"Launch Server"\nbutton, Config panel',
            'actions': 'Selects H200 config,\nclicks "Launch Server"',
            'emotions': 'Excitement',
            'pain': 'Queue if pool\nis fully allocated',
            'color': BRAND_GREEN
        },
        {
            'name': '5. COMPUTE',
            'emoji': '>>',
            'touchpoints': 'SSH endpoint,\nJupyter, API',
            'actions': 'Trains models,\nruns inference',
            'emotions': 'Delight',
            'pain': 'Session time\nmanagement',
            'color': BRAND_ACCENT
        },
        {
            'name': '6. RETAIN',
            'emoji': '>>',
            'touchpoints': 'Usage dashboard,\nBilling history',
            'actions': 'Reviews usage,\ntops up again',
            'emotions': 'Loyalty',
            'pain': 'Scaling needs\nbeyond 4 GPU',
            'color': '#D32F2F'
        },
    ]
    
    stage_width = 2.0
    gap = 0.2
    start_x = 0.4
    
    # Draw connection line
    ax.plot([start_x + stage_width/2, start_x + (len(stages)-1)*(stage_width+gap) + stage_width/2],
            [5.2, 5.2], color='#DDD', linewidth=3, zorder=1)
    
    # Emotion curve
    emotion_values = [3, 5, 6, 7.5, 9, 8.5]
    emotion_x = [start_x + i*(stage_width+gap) + stage_width/2 for i in range(len(stages))]
    emotion_y_base = 1.8
    emotion_scale = 0.25
    emotion_y = [emotion_y_base + v * emotion_scale for v in emotion_values]
    
    # Fill under curve
    ax.fill_between(emotion_x, emotion_y_base, emotion_y,
                    alpha=0.15, color=BRAND_GREEN, zorder=1)
    ax.plot(emotion_x, emotion_y, color=BRAND_GREEN, linewidth=2.5,
            marker='o', markersize=8, markerfacecolor='white',
            markeredgecolor=BRAND_GREEN, markeredgewidth=2, zorder=2)
    
    ax.text(0.1, emotion_y_base - 0.25, 'EMOTIONAL CURVE',
            fontsize=8, fontweight='bold', color=BRAND_GREEN, alpha=0.7)
    
    for i, stage in enumerate(stages):
        x = start_x + i * (stage_width + gap)
        color = stage['color']
        
        # Stage header circle
        circle = plt.Circle((x + stage_width/2, 5.2), 0.35,
                           facecolor=color, edgecolor='white',
                           linewidth=2, zorder=3)
        ax.add_patch(circle)
        ax.text(x + stage_width/2, 5.2, stage['emoji'],
                ha='center', va='center', fontsize=14, zorder=4)
        
        # Stage name
        ax.text(x + stage_width/2, 5.85, stage['name'],
                ha='center', va='bottom', fontsize=7.5,
                fontweight='bold', color=color)
        
        # Touchpoints card
        card = FancyBboxPatch(
            (x + 0.05, 3.7), stage_width - 0.1, 1.1,
            boxstyle="round,pad=0.08",
            facecolor='white', edgecolor=color,
            linewidth=1.5, alpha=0.9, zorder=2
        )
        ax.add_patch(card)
        ax.text(x + stage_width/2, 4.6, 'Touchpoints',
                ha='center', fontsize=6.5, fontweight='bold', color=color)
        ax.text(x + stage_width/2, 4.2, stage['touchpoints'],
                ha='center', va='center', fontsize=6, color='#555',
                linespacing=1.2)
        
        # Actions card
        card2 = FancyBboxPatch(
            (x + 0.05, 2.4), stage_width - 0.1, 1.1,
            boxstyle="round,pad=0.08",
            facecolor=color, edgecolor=color,
            linewidth=1, alpha=0.12, zorder=2
        )
        ax.add_patch(card2)
        ax.text(x + stage_width/2, 3.3, 'Actions',
                ha='center', fontsize=6.5, fontweight='bold', color=color)
        ax.text(x + stage_width/2, 2.9, stage['actions'],
                ha='center', va='center', fontsize=6, color='#444',
                linespacing=1.2)
        
        # Emotion label
        ax.text(x + stage_width/2, emotion_y[i] + 0.15, stage['emotions'],
                ha='center', fontsize=7, color=color, fontweight='bold')
    
    # Legend
    ax.text(7, 0.5,
            'Pain Points: Awareness gap -> GPU pricing -> Wallet conversion -> Pool queues -> Session mgmt -> Scaling',
            ha='center', fontsize=8, color='#666', style='italic')
    
    plt.tight_layout()
    path = os.path.join(CHARTS_DIR, 'cjm_diagram.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def create_flow_diagram():
    """Process flow diagram — updated for Zinda IB UI -> API -> zypl.ai provisioning."""
    fig, ax = plt.subplots(figsize=(14, 10))
    fig.patch.set_facecolor('#FAFBFE')
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')
    
    ax.text(7, 9.5, 'SYSTEM FLOW — VIRTUAL COMPUTE WALLET',
            ha='center', fontsize=15, fontweight='bold', color=BRAND_PRIMARY)
    ax.text(7, 9.15, 'Zinda Internet Banking UI  ->  API  ->  zypl.ai Provisioning',
            ha='center', fontsize=10, fontstyle='italic', color='#777')
    
    # Swimlane headers
    lanes = [
        ('USER\n(Zinda IB App)', BRAND_BLUE, 8.0),
        ('ZINDA CAPITAL\n(Ledger + API)', BRAND_ACCENT, 5.0),
        ('ZYPL.AI\n(GPU Infra)', BRAND_GREEN, 2.0),
    ]
    
    for name, color, y in lanes:
        # Lane background
        lane_rect = FancyBboxPatch(
            (0.3, y - 1.2), 13.4, 2.6,
            boxstyle="round,pad=0.1",
            facecolor=color, edgecolor=color,
            linewidth=0, alpha=0.06
        )
        ax.add_patch(lane_rect)
        
        # Lane label
        label_rect = FancyBboxPatch(
            (0.3, y - 0.3), 2.0, 0.6,
            boxstyle="round,pad=0.05",
            facecolor=color, edgecolor=color,
            linewidth=1, alpha=0.9
        )
        ax.add_patch(label_rect)
        ax.text(1.3, y, name, ha='center', va='center',
                fontsize=7, fontweight='bold', color='white', linespacing=1.1)
    
    # Process boxes function
    def draw_process_box(x, y, w, h, text, color, shape='rect'):
        if shape == 'diamond':
            diamond = plt.Polygon(
                [(x, y), (x + w/2, y + h/2), (x + w, y), (x + w/2, y - h/2)],
                facecolor='white', edgecolor=color, linewidth=2
            )
            ax.add_patch(diamond)
            ax.text(x + w/2, y, text, ha='center', va='center',
                    fontsize=6.5, fontweight='bold', color=color, linespacing=1.1)
        elif shape == 'circle':
            circle = plt.Circle((x + w/2, y), w/2,
                              facecolor=color, edgecolor='white',
                              linewidth=2, alpha=0.9)
            ax.add_patch(circle)
            ax.text(x + w/2, y, text, ha='center', va='center',
                    fontsize=6.5, fontweight='bold', color='white', linespacing=1.1)
        else:
            rect = FancyBboxPatch(
                (x, y - h/2), w, h,
                boxstyle="round,pad=0.08",
                facecolor='white', edgecolor=color,
                linewidth=2, alpha=0.95
            )
            ax.add_patch(rect)
            ax.text(x + w/2, y, text, ha='center', va='center',
                    fontsize=6.5, fontweight='bold', color=color, linespacing=1.1)
    
    def draw_arrow(x1, y1, x2, y2, color='#999', label=''):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=color,
                                   lw=1.8, connectionstyle='arc3,rad=0'))
        if label:
            mid_x = (x1 + x2) / 2
            mid_y = (y1 + y2) / 2
            ax.text(mid_x, mid_y + 0.15, label, ha='center', fontsize=5.5,
                    color=color, fontstyle='italic')
    
    # ── USER lane (y=8.0) ──
    draw_process_box(2.8, 8.0, 1.2, 0.65, 'START\nOpen App', BRAND_BLUE, 'circle')
    draw_process_box(4.5, 8.0, 1.8, 0.65, 'Navigate to\n"Cloud Compute"\nmodule', BRAND_BLUE)
    draw_process_box(6.8, 8.0, 1.8, 0.65, 'Accept T&C\n& Fund Wallet\n(TJS)', BRAND_BLUE)
    draw_process_box(9.1, 8.0, 1.8, 0.65, 'Click\n"Launch Server"\n(H200)', BRAND_BLUE)
    draw_process_box(11.4, 8.0, 1.8, 0.65, 'Receive\nEndpoint +\nCredentials', BRAND_BLUE)
    
    draw_arrow(3.95, 8.0, 4.5, 8.0, BRAND_BLUE)
    draw_arrow(6.3, 8.0, 6.8, 8.0, BRAND_BLUE)
    draw_arrow(8.6, 8.0, 9.1, 8.0, BRAND_BLUE)
    draw_arrow(10.9, 8.0, 11.4, 8.0, BRAND_BLUE)
    
    # ── ZINDA lane (y=5.0) ──
    draw_process_box(3.5, 5.0, 1.8, 0.65, 'Receive\n"Launch"\nAPI request', BRAND_ACCENT)
    draw_process_box(5.8, 5.0, 1.8, 0.65, 'Check wallet\nbalance /\ncredit limit', BRAND_ACCENT)
    draw_process_box(8.1, 5.0, 1.6, 0.8, 'Sufficient\nfunds?', BRAND_ACCENT, 'diamond')
    draw_process_box(10.2, 5.0, 1.8, 0.65, 'Call zypl.ai\nprovisioning\nendpoint', BRAND_ACCENT)
    draw_process_box(12.3, 5.0, 1.4, 0.65, 'Debit wallet\nper metering\nfeed', BRAND_ACCENT)
    
    draw_arrow(5.3, 5.0, 5.8, 5.0, BRAND_ACCENT)
    draw_arrow(7.6, 5.0, 8.1, 5.0, BRAND_ACCENT)
    draw_arrow(9.7, 5.0, 10.2, 5.0, BRAND_ACCENT)
    draw_arrow(12.0, 5.0, 12.3, 5.0, BRAND_ACCENT)
    
    # ── ZYPL lane (y=2.0) ──
    draw_process_box(3.5, 2.0, 1.8, 0.65, 'Receive\nprovisioning\nrequest', BRAND_GREEN)
    draw_process_box(5.8, 2.0, 1.8, 0.65, 'Allocate\n1x H200\nfrom pool', BRAND_GREEN)
    draw_process_box(8.1, 2.0, 1.8, 0.65, 'Provision\ninstance &\ngen creds', BRAND_GREEN)
    draw_process_box(10.4, 2.0, 1.8, 0.65, 'Real-time\nGPU-hour\nmetering', BRAND_GREEN)
    draw_process_box(12.5, 2.0, 1.2, 0.65, 'Usage\nreport\n-> Zinda', BRAND_GREEN, 'circle')
    
    draw_arrow(5.3, 2.0, 5.8, 2.0, BRAND_GREEN)
    draw_arrow(7.6, 2.0, 8.1, 2.0, BRAND_GREEN)
    draw_arrow(9.9, 2.0, 10.4, 2.0, BRAND_GREEN)
    draw_arrow(12.2, 2.0, 12.5, 2.0, BRAND_GREEN)
    
    # Cross-lane arrows
    draw_arrow(10.0, 7.65, 4.4, 5.4, '#999')       # User Launch -> Zinda receive
    draw_arrow(11.1, 4.6, 4.4, 2.4, '#999')         # Zinda call -> zypl receive
    draw_arrow(9.0, 2.4, 12.0, 7.65, '#999')        # zypl creds -> user endpoint
    draw_arrow(12.8, 2.6, 13.0, 4.6, '#999')        # zypl report -> zinda billing
    
    # Legend at bottom
    legend_items = [
        ('User actions (Zinda IB App)', BRAND_BLUE),
        ('Zinda Capital (Ledger + API)', BRAND_ACCENT),
        ('zypl.ai (GPU Infrastructure)', BRAND_GREEN),
    ]
    for i, (label, color) in enumerate(legend_items):
        x = 2.5 + i * 3.8
        rect = FancyBboxPatch(
            (x, 0.3), 0.4, 0.3,
            boxstyle="round,pad=0.03",
            facecolor=color, edgecolor=color,
            linewidth=1, alpha=0.8
        )
        ax.add_patch(rect)
        ax.text(x + 0.55, 0.45, label, fontsize=8, color='#555', va='center')
    
    plt.tight_layout()
    path = os.path.join(CHARTS_DIR, 'flow_diagram.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def create_market_comparison_chart():
    """Market comparison chart showing Virtual Compute Wallet vs alternatives."""
    fig, ax = plt.subplots(figsize=(10, 5))
    fig.patch.set_facecolor('#FAFBFE')
    
    providers = ['Virtual Compute\nWallet (zypl.ai)', 'AWS\np4d.24xlarge', 'Google Cloud\nA3-mega', 'Azure\nND H200 v5', 'Lambda Cloud\nH200']
    prices = [3.5, 32.77, 34.72, 28.43, 3.99]
    colors = [BRAND_ACCENT, '#FF9900', '#4285F4', '#0078D4', '#8B5CF6']
    
    bars = ax.barh(providers, prices, color=colors, alpha=0.85,
                   edgecolor='white', linewidth=1, height=0.6)
    
    for bar, price in zip(bars, prices):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                f'${price:.2f}/hr', va='center', fontsize=10,
                fontweight='bold', color='#333')
    
    ax.set_xlabel('Cost ($/GPU-hour)', fontsize=11, fontweight='bold', color='#333')
    ax.set_title('GPU Compute Cost Comparison',
                 fontsize=14, fontweight='bold', color=BRAND_PRIMARY, pad=15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_xlim(0, 42)
    ax.grid(axis='x', alpha=0.3, linestyle='--')
    
    # Add "best value" badge
    ax.annotate('BEST VALUE ★', xy=(3.5, 0),
                xytext=(8, -0.3), fontsize=9, fontweight='bold',
                color=BRAND_ACCENT,
                arrowprops=dict(arrowstyle='->', color=BRAND_ACCENT, lw=1.5))
    
    plt.tight_layout()
    path = os.path.join(CHARTS_DIR, 'market_comparison.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def create_model_comparison_chart():
    """Prepaid Wallet vs Credit Limit model comparison."""
    fig, axes = plt.subplots(1, 2, figsize=(12, 5))
    fig.patch.set_facecolor('#FAFBFE')
    fig.suptitle('Model Comparison: Prepaid Wallet vs Credit Limit',
                 fontsize=13, fontweight='bold', color=BRAND_PRIMARY, y=1.02)
    
    # Model A - Prepaid Wallet
    ax1 = axes[0]
    criteria = ['Launch\nSpeed', 'Credit\nRisk', 'Compliance', 'User\nAttraction', 'Scala-\nbility']
    scores_a = [9, 9, 8, 6, 7]
    scores_b = [5, 3, 4, 9, 8]
    
    angles = np.linspace(0, 2*np.pi, len(criteria), endpoint=False).tolist()
    scores_a_plot = scores_a + [scores_a[0]]
    scores_b_plot = scores_b + [scores_b[0]]
    angles += [angles[0]]
    
    ax1.set_ylim(0, 10)
    ax1 = fig.add_subplot(121, polar=True)
    ax1.fill(angles, scores_a_plot, alpha=0.25, color=BRAND_GREEN)
    ax1.plot(angles, scores_a_plot, 'o-', color=BRAND_GREEN, linewidth=2,
             label='A: Prepaid Wallet', markersize=6)
    ax1.fill(angles, scores_b_plot, alpha=0.15, color=BRAND_ACCENT)
    ax1.plot(angles, scores_b_plot, 's--', color=BRAND_ACCENT, linewidth=2,
             label='B: Credit Limit', markersize=6)
    ax1.set_xticks(angles[:-1])
    ax1.set_xticklabels(criteria, fontsize=8)
    ax1.set_ylim(0, 10)
    ax1.set_title('Radar Comparison\nof Models', fontsize=11,
                  fontweight='bold', color='#333', pad=20)
    ax1.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize=9)
    
    # Model comparison bar chart
    ax2 = axes[1]
    categories = ['Time to\nLaunch', 'Bank\nRisk', 'Regulatory\nComplexity', 'Entry\nBarrier', 'User\nAppeal']
    val_a = [9, 2, 3, 7, 6]
    val_b = [5, 8, 7, 3, 9]
    
    x = np.arange(len(categories))
    width = 0.35
    ax2.bar(x - width/2, val_a, width, label='A: Prepaid Wallet',
            color=BRAND_GREEN, alpha=0.8, edgecolor='white')
    ax2.bar(x + width/2, val_b, width, label='B: Credit Limit',
            color=BRAND_ACCENT, alpha=0.8, edgecolor='white')
    ax2.set_xticks(x)
    ax2.set_xticklabels(categories, fontsize=8)
    ax2.set_ylabel('Score (1-10)', fontsize=10)
    ax2.set_title('Component Comparison', fontsize=11,
                  fontweight='bold', color='#333')
    ax2.legend(fontsize=9)
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.set_ylim(0, 11)
    
    plt.tight_layout()
    path = os.path.join(CHARTS_DIR, 'model_comparison.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def create_architecture_diagram():
    """System architecture diagram for the Virtual Compute Wallet."""
    fig, ax = plt.subplots(figsize=(14, 8))
    fig.patch.set_facecolor('#FAFBFE')
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')

    ax.text(7, 7.6, 'SYSTEM ARCHITECTURE — VIRTUAL COMPUTE WALLET',
            ha='center', fontsize=15, fontweight='bold', color=BRAND_PRIMARY)
    ax.text(7, 7.25, 'Embedded Module in Zinda Internet Banking Platform',
            ha='center', fontsize=10, fontstyle='italic', color='#777')

    # ── Zinda IB App (top block) ──
    app_rect = FancyBboxPatch(
        (1.0, 4.5), 12.0, 2.4,
        boxstyle="round,pad=0.15",
        facecolor='#EBF5FB', edgecolor=BRAND_BLUE,
        linewidth=2.5, alpha=0.95
    )
    ax.add_patch(app_rect)
    ax.text(7, 6.7, 'ZINDA INTERNET BANKING APP',
            ha='center', fontsize=12, fontweight='bold', color=BRAND_BLUE)

    # Sub-modules inside the app
    modules = [
        (1.5, 4.8, 2.5, 1.4, '"Cloud Compute"\nModule', BRAND_PURPLE,
         ['- Wallet balance', '- Top-up flow', '- Launch button']),
        (4.5, 4.8, 2.5, 1.4, 'Server Config\nPanel', BRAND_TEAL,
         ['- GPU selection', '- Instance size', '- Cost estimate']),
        (7.5, 4.8, 2.5, 1.4, 'Usage\nDashboard', BRAND_GREEN,
         ['- Real-time usage', '- GPU-hours log', '- Billing history']),
        (10.5, 4.8, 2.2, 1.4, 'Credentials\nDelivery', BRAND_GOLD,
         ['- SSH endpoint', '- API key', '- Jupyter URL']),
    ]

    for (x, y, w, h, title, color, items) in modules:
        mod_rect = FancyBboxPatch(
            (x, y), w, h,
            boxstyle="round,pad=0.08",
            facecolor='white', edgecolor=color,
            linewidth=1.8, alpha=0.95
        )
        ax.add_patch(mod_rect)
        ax.text(x + w/2, y + h - 0.25, title,
                ha='center', va='center', fontsize=7.5,
                fontweight='bold', color=color, linespacing=1.1)
        content_y = y + h - 0.65
        for item in items:
            ax.text(x + 0.15, content_y, item,
                    ha='left', fontsize=6, color='#555')
            content_y -= 0.25

    # ── Backend Services (bottom left) ──
    be1_rect = FancyBboxPatch(
        (1.0, 0.8), 5.5, 3.0,
        boxstyle="round,pad=0.15",
        facecolor='#FDEDEC', edgecolor=BRAND_ACCENT,
        linewidth=2.5, alpha=0.95
    )
    ax.add_patch(be1_rect)
    ax.text(3.75, 3.55, 'ZINDA BACKEND (Ledger)',
            ha='center', fontsize=11, fontweight='bold', color=BRAND_ACCENT)

    zinda_items = [
        ('TJS Wallet Account', '- Balance tracking\n- Transaction log'),
        ('Billing Engine', '- Metering ingest\n- Statement gen'),
        ('KYC / Compliance', '- User verification\n- Regulatory reports'),
    ]
    for i, (title, desc) in enumerate(zinda_items):
        bx = 1.4 + i * 1.8
        item_rect = FancyBboxPatch(
            (bx, 1.0), 1.6, 2.0,
            boxstyle="round,pad=0.06",
            facecolor='white', edgecolor=BRAND_ACCENT,
            linewidth=1, alpha=0.9
        )
        ax.add_patch(item_rect)
        ax.text(bx + 0.8, 2.75, title, ha='center', fontsize=6.5,
                fontweight='bold', color=BRAND_ACCENT)
        ax.text(bx + 0.8, 2.15, desc, ha='center', va='center',
                fontsize=5.5, color='#555', linespacing=1.3)

    # ── zypl.ai Backend (bottom right) ──
    be2_rect = FancyBboxPatch(
        (7.5, 0.8), 5.5, 3.0,
        boxstyle="round,pad=0.15",
        facecolor='#E8F8F5', edgecolor=BRAND_GREEN,
        linewidth=2.5, alpha=0.95
    )
    ax.add_patch(be2_rect)
    ax.text(10.25, 3.55, 'ZYPL.AI BACKEND (GPU Infra)',
            ha='center', fontsize=11, fontweight='bold', color=BRAND_GREEN)

    zypl_items = [
        ('Provisioning API', '- /compute/launch\n- /compute/terminate'),
        ('GPU Pool Mgr', '- 4x H200 alloc\n- Queue mgmt'),
        ('Metering Engine', '- Real-time tracking\n- Usage reports'),
    ]
    for i, (title, desc) in enumerate(zypl_items):
        bx = 7.9 + i * 1.8
        item_rect = FancyBboxPatch(
            (bx, 1.0), 1.6, 2.0,
            boxstyle="round,pad=0.06",
            facecolor='white', edgecolor=BRAND_GREEN,
            linewidth=1, alpha=0.9
        )
        ax.add_patch(item_rect)
        ax.text(bx + 0.8, 2.75, title, ha='center', fontsize=6.5,
                fontweight='bold', color=BRAND_GREEN)
        ax.text(bx + 0.8, 2.15, desc, ha='center', va='center',
                fontsize=5.5, color='#555', linespacing=1.3)

    # ── Arrows between layers ──
    # App -> Zinda Backend
    ax.annotate('', xy=(3.75, 3.85), xytext=(3.75, 4.5),
                arrowprops=dict(arrowstyle='->', color=BRAND_ACCENT, lw=2.5))
    ax.text(3.75, 4.15, 'API calls', ha='center', fontsize=7, color=BRAND_ACCENT, fontweight='bold')

    # App -> zypl Backend
    ax.annotate('', xy=(10.25, 3.85), xytext=(10.25, 4.5),
                arrowprops=dict(arrowstyle='->', color=BRAND_GREEN, lw=2.5))
    ax.text(10.25, 4.15, 'Provisioning', ha='center', fontsize=7, color=BRAND_GREEN, fontweight='bold')

    # Zinda <-> zypl
    ax.annotate('', xy=(7.5, 2.3), xytext=(6.5, 2.3),
                arrowprops=dict(arrowstyle='<->', color='#999', lw=2))
    ax.text(7.0, 2.55, 'REST API\n(metering, billing)', ha='center', fontsize=6.5,
            color='#777', fontweight='bold', linespacing=1.2)

    plt.tight_layout()
    path = os.path.join(CHARTS_DIR, 'architecture_diagram.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


# ═══════════════════════════════════════════════════════════════
#  MAIN DOCUMENT GENERATION
# ═══════════════════════════════════════════════════════════════

def generate_document():
    """Generate the full professional document — v2.0 Virtual Compute Wallet."""
    print("[*] Generating charts and visualizations...")
    
    # Generate all charts
    revenue_chart = create_revenue_projection_chart()
    print("  [OK] Revenue projection")
    
    gpu_metrics = create_gpu_utilization_gauge()
    print("  [OK] GPU metrics")
    
    bmc_chart = create_business_model_canvas()
    print("  [OK] Business model canvas")
    
    cjm_chart = create_cjm_diagram()
    print("  [OK] CJM diagram")
    
    flow_chart = create_flow_diagram()
    print("  [OK] Flow diagram")
    
    market_chart = create_market_comparison_chart()
    print("  [OK] Market comparison")
    
    model_comp = create_model_comparison_chart()
    print("  [OK] Model comparison")
    
    arch_diagram = create_architecture_diagram()
    print("  [OK] Architecture diagram")
    
    print("\n[*] Forming Word document...")
    
    doc = Document()
    
    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    
    # ═══════════════════════════════════════════════════════════
    #  TITLE PAGE
    # ═══════════════════════════════════════════════════════════
    
    for _ in range(4):
        doc.add_paragraph('')
    
    p = add_formatted_paragraph(
        doc, 'VIRTUAL COMPUTE WALLET', font_name='Calibri', size=36,
        bold=True, color=BRAND_PRIMARY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )
    
    add_formatted_paragraph(
        doc, 'Embedded GPU Credit Module', font_name='Calibri', size=18,
        bold=True, color=BRAND_PURPLE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8
    )
    
    add_formatted_paragraph(
        doc, 'zypl.ai \u00d7 Zinda Capital', font_name='Calibri', size=20,
        bold=True, color=BRAND_ACCENT,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20
    )
    
    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2501' * 40)
    run.font.color.rgb = RGBColor(*hex_to_rgb(BRAND_ACCENT))
    run.font.size = Pt(14)
    
    add_formatted_paragraph(
        doc, 'On-demand GPU compute access via Zinda Internet Banking', font_name='Calibri', size=16,
        italic=True, color='#666666',
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=4
    )
    
    add_formatted_paragraph(
        doc, 'NVIDIA H200 \u2022 TJS Wallet \u2022 Prepaid & Credit Models',
        font_name='Calibri', size=12, italic=True, color='#888888',
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30
    )
    
    # Document info box
    info_lines = [
        'PRODUCT SPECIFICATIONS  \u2022  CUSTOMER JOURNEY MAP (CJM)',
        'BUSINESS MODEL  \u2022  SYSTEM ARCHITECTURE  \u2022  FLOW DIAGRAM',
        '',
        'Document version: v2.0  |  June 2026',
        'Post-meeting update: Akai Akmal (Visa) feedback incorporated',
        'Confidential'
    ]
    for line in info_lines:
        add_formatted_paragraph(
            doc, line, font_name='Calibri',
            size=11 if '\u2022' in line else 9,
            bold='\u2022' in line,
            color=BRAND_PRIMARY if '\u2022' in line else '#999999',
            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3
        )
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════
    
    add_section_title(doc, 'TABLE OF CONTENTS', 1)
    
    toc_items = [
        ('1.', 'Product Specifications', ''),
        ('  1.1', 'Product Identity (v1.0 vs v2.0)', ''),
        ('  1.2', 'Value Proposition', ''),
        ('  1.3', 'Technical Specifications', ''),
        ('  1.4', 'Key Infrastructure Metrics', ''),
        ('  1.5', 'Participants & Roles', ''),
        ('2.', 'Product Architecture', ''),
        ('  2.1', 'System Architecture Diagram', ''),
        ('  2.2', 'Backend: TJS Wallet Account', ''),
        ('  2.3', 'Frontend: Visual Compute Interface', ''),
        ('  2.4', 'API Integration Points', ''),
        ('3.', 'Dual Product Models', ''),
        ('  3.1', 'Model A: Prepaid Wallet (Recommended)', ''),
        ('  3.2', 'Model B: Credit Limit', ''),
        ('  3.3', 'Model Comparison', ''),
        ('4.', 'Customer Journey Map (CJM)', ''),
        ('  4.1', 'Visual Journey Map', ''),
        ('  4.2', 'Journey Stage Details', ''),
        ('  4.3', 'v1.0 vs v2.0 CJM Changes', ''),
        ('  4.4', 'Pain Points & Solutions', ''),
        ('5.', 'Business Model', ''),
        ('  5.1', 'Business Model Canvas', ''),
        ('  5.2', 'Financial Projections', ''),
        ('  5.3', 'Competitive Analysis', ''),
        ('6.', 'System Flow Diagram', ''),
        ('  6.1', 'Process Flow (Swimlane)', ''),
        ('  6.2', 'Process Step Descriptions', ''),
        ('7.', 'Advantages of New Architecture', ''),
        ('8.', 'Implementation Roadmap', ''),
        ('9.', 'Open Questions & Next Steps', ''),
    ]
    
    for num, title, page in toc_items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(2)
        
        run_num = p.add_run(f'{num}  ')
        run_num.font.name = 'Calibri'
        run_num.font.size = Pt(10)
        run_num.font.bold = True
        run_num.font.color.rgb = RGBColor(*hex_to_rgb(BRAND_ACCENT))
        
        is_sub = num.startswith(' ')
        run_title = p.add_run(title)
        run_title.font.name = 'Calibri'
        run_title.font.size = Pt(10 if not is_sub else 9)
        run_title.font.bold = not is_sub
        run_title.font.color.rgb = RGBColor(*hex_to_rgb(BRAND_PRIMARY if not is_sub else '#666666'))
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════
    #  PIVOT NOTICE
    # ═══════════════════════════════════════════════════════════
    
    add_section_title(doc, 'IMPORTANT: PRODUCT PIVOT NOTICE', 1)
    
    add_formatted_paragraph(
        doc,
        'Following the meeting with Akai Akmal (Visa representative) on June 2, 2026, '
        'the product has been pivoted from a physical co-branded Visa card to an embedded '
        'digital module within Zinda\'s Internet Banking platform.',
        size=10, bold=True, color=BRAND_ACCENT, space_after=8
    )
    
    pivot_reasons = [
        ['Constraint', 'Detail'],
        ['Visa BIN requirement', 'Visa requires 10,000+ cards to issue a dedicated BIN — impractical for a niche GPU compute pilot'],
        ['Per-card compliance', 'Each printed card requires specific Visa compliance certification'],
        ['Recommendation', 'Akai Akmal recommended embedding the product inside Zinda Internet Banking as an "Embedded GPU Credit Module" or "Virtual Compute Wallet"'],
        ['New form factor', 'TJS wallet with a front-end visual and back-end on a wallet account — no physical card'],
    ]
    
    add_styled_table(doc, pivot_reasons[0], pivot_reasons[1:],
                     header_color=BRAND_ACCENT,
                     col_widths=[4, 12])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════
    #  SECTION 1: PRODUCT SPECIFICATIONS
    # ═══════════════════════════════════════════════════════════
    
    add_section_title(doc, '1. PRODUCT SPECIFICATIONS', 1)
    
    add_section_title(doc, '1.1 Product Identity — v1.0 vs v2.0', 2)
    
    identity_data = [
        ['Parameter', 'v1.0 (Old)', 'v2.0 (Current)'],
        ['Product Name', 'GPU Credit Card', 'Virtual Compute Wallet'],
        ['Working Title', 'GPU Credit Card', '"Cloud Compute" module'],
        ['Product Type', 'Physical co-branded Visa card', 'Embedded module in Zinda Internet Banking'],
        ['Form Factor', 'Physical plastic card', 'Digital wallet with front-end UI'],
        ['Backend', 'Card payment rail + Visa BIN', 'TJS wallet account on Zinda\'s ledger'],
        ['Card Issuance', 'Zinda issues physical card', 'No card — user activates module in-app'],
        ['Visa Dependency', 'Required (BIN, co-branding)', 'None — fully in-house via Zinda'],
        ['Regulatory Path', 'Card product regulation', 'E-wallet / prepaid instrument regulation'],
        ['Market', 'Tajikistan', 'Tajikistan'],
        ['Infrastructure Provider', 'zypl.ai (GPU fleet)', 'zypl.ai (GPU fleet + scoring)'],
        ['Target Hardware', 'NVIDIA H200 (141 GB HBM3e)', 'NVIDIA H200 (141 GB HBM3e)'],
        ['Payment Model', 'Pay-per-use GPU-hours', 'Prepaid wallet OR credit limit'],
        ['Currencies', 'TJS / USD', 'TJS / USD'],
    ]
    
    add_styled_table(doc, identity_data[0], identity_data[1:],
                     header_color=BRAND_PRIMARY,
                     col_widths=[3.5, 5.5, 7])
    
    add_formatted_paragraph(doc, '', size=6)
    
    # Value Proposition
    add_section_title(doc, '1.2 Value Proposition', 2)
    
    value_props = [
        ('\U0001F3AF', 'Problem', 'Tajikistan has no domestic GPU-on-demand offering. The Darya cluster serves only zypl.ai and UZCARD. Researchers and startups must use foreign clouds with premium pricing, FX friction, and data sovereignty risks.'),
        ('\U0001F4A1', 'Solution', 'An embedded module inside Zinda Internet Banking provides instant access to NVIDIA H200 GPU compute. Users fund a TJS wallet and click "Launch Server" — no physical card, no FX, no foreign account.'),
        ('\U0001F3C6', 'Competitive Advantage', 'Zero domestic competition. Sovereign compute rail. Pricing 8-10x lower than major clouds (AWS, GCP, Azure). Leverages Zinda\'s existing user base — no new KYC needed.'),
        ('\U0001F680', 'Scaling Strategy', 'Pilot with 4 GPUs validates the rail and locks demand. Linear scaling: 100 GPUs = $2.25M-$4.25M/year revenue.'),
    ]
    
    for emoji, title, desc in value_props:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(6)
        pf.left_indent = Cm(0.5)
        
        run = p.add_run(f'{emoji}  {title}: ')
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*hex_to_rgb(BRAND_ACCENT))
        
        run2 = p.add_run(desc)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(*hex_to_rgb('#444444'))
    
    add_section_title(doc, '1.3 Technical Specifications', 2)
    
    tech_spec = [
        ['Parameter', 'Value', 'Notes'],
        ['GPU Model', 'NVIDIA H200 SXM', 'Flagship AI GPU, Hopper architecture'],
        ['GPU Memory', '141 GB HBM3e', 'Largest available on market'],
        ['Pilot Pool', '4\u00d7 H200', 'Single physical server'],
        ['Unit of Allocation', '1\u00d7 H200 per session', 'Dedicated access'],
        ['Max Concurrent Sessions', '4', 'Limited by pool size'],
        ['Total GPU-hours/month', '2,880 hrs', '4 GPUs \u00d7 720 hours'],
        ['Market Rate', '$3-4 / GPU-hour', 'Competitive with Lambda Cloud'],
        ['Billing', 'Real-time metering', 'zypl.ai API \u2192 Zinda ledger'],
        ['Wallet Currency', 'TJS (Tajikistani Somoni)', 'With USD display equivalent'],
    ]
    
    add_styled_table(doc, tech_spec[0], tech_spec[1:],
                     header_color=BRAND_BLUE,
                     col_widths=[4.5, 4.5, 7])
    
    add_formatted_paragraph(doc, '', size=4)
    
    # GPU Metrics Chart
    add_section_title(doc, '1.4 Key Infrastructure Metrics', 2)
    doc.add_picture(gpu_metrics, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_formatted_paragraph(doc, '', size=4)
    
    # Roles table
    add_section_title(doc, '1.5 Participants & Roles', 2)
    
    roles_data = [
        ['Participant', 'Role', 'Responsibilities'],
        ['Zinda Capital', 'Platform owner,\nwallet issuer',
         'Internet banking UI, wallet management, KYC, billing, credit underwriting, regulatory compliance'],
        ['zypl.ai', 'Infrastructure provider,\ncredit scoring engine',
         'GPU fleet operations (4\u00d7 H200), instance provisioning, GPU-hour metering, scoring API, usage reporting.'],
        ['End User\n(Individual)', 'Consumer\n(Mobile Wallet)',
         'Researchers, professors, university teachers of AI \u2014 accesses via dedicated mobile wallet app. '
         'Wallet is purpose-locked: funds can ONLY be used for GPU compute.'],
        ['End User\n(Company)', 'Consumer\n(Internet Banking)',
         'AI startups, tech companies, financial companies, fintechs \u2014 accesses via Zinda Internet Banking. '
         'Dedicated e-wallet account purpose-locked for GPU compute only.'],
        ['Central Bank\nof Tajikistan', 'Regulator',
         'E-wallet/prepaid instrument oversight, credit product licensing (for v2)'],
    ]
    
    add_styled_table(doc, roles_data[0], roles_data[1:],
                     header_color=BRAND_PURPLE,
                     col_widths=[3, 3.5, 9.5])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════
    #  SECTION 2: PRODUCT ARCHITECTURE
    # ═══════════════════════════════════════════════════════════
    
    add_section_title(doc, '2. PRODUCT ARCHITECTURE', 1)
    
    add_formatted_paragraph(
        doc,
        'The Virtual Compute Wallet is implemented as an embedded module within Zinda\'s '
        'Internet Banking platform. The frontend provides a visual compute interface; the '
        'backend runs on a TJS wallet account within Zinda\'s ledger, integrated with '
        'zypl.ai\'s GPU provisioning API.',
        size=10, color='#555555', space_after=12
    )
    
    add_section_title(doc, '2.1 System Architecture Diagram', 2)
    doc.add_picture(arch_diagram, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_formatted_paragraph(doc, '', size=6)
    
    add_section_title(doc, '2.2 Backend: TJS Wallet Account', 2)
    
    wallet_data = [
        ['Component', 'Description'],
        ['Account Type', 'Dedicated e-wallet / internet banking account (NOT a general-purpose account)'],
        ['Purpose Lock', 'Funds are PURPOSE-LOCKED: can ONLY be used for GPU compute services. '
         'Cannot be transferred, withdrawn, or applied to any other purpose.'],
        ['Currency', 'TJS (Tajikistani Somoni), with USD display equivalent'],
        ['Ledger', 'Zinda\'s core banking ledger'],
        ['Balance Types', 'Prepaid balance (user-funded) OR Credit balance (Zinda-issued limit, scored by zypl.ai)'],
        ['Transaction Model', 'Debit from wallet \u2192 credit to zypl.ai per GPU-hour consumed'],
        ['Reconciliation', 'Real-time metering feed from zypl.ai API \u2192 Zinda ledger entries'],
        ['Channel: Mobile Wallet', 'For individual researchers, university professors, teachers of AI \u2014 '
         'dedicated mobile wallet app with purpose-locked GPU compute funds'],
        ['Channel: Internet Banking', 'For AI startups, tech companies, financial companies, fintechs \u2014 '
         'dedicated e-wallet account within Zinda Internet Banking, purpose-locked for GPU compute'],
    ]
    
    add_styled_table(doc, wallet_data[0], wallet_data[1:],
                     header_color=BRAND_PRIMARY,
                     col_widths=[4, 12])
    
    add_formatted_paragraph(doc, '', size=6)
    
    add_section_title(doc, '2.3 Frontend: Visual Compute Interface', 2)
    
    frontend_data = [
        ['UI Component', 'Description'],
        ['Module Entry', '"Cloud Compute" tab/section in the main Zinda app navigation'],
        ['Wallet View', 'Shows current TJS balance, available GPU-hours equivalent, usage history'],
        ['Server Launcher', '"Launch Server" button that triggers backend API call to zypl.ai'],
        ['Configuration Panel', 'GPU type selection (H200), instance size, estimated cost display'],
        ['Usage Dashboard', 'Real-time GPU utilization, cost accumulation, session history'],
        ['Top-up Flow', 'Add funds from linked bank account or card'],
        ['Terms & Conditions', 'One-time acceptance screen on first access to the module'],
    ]
    
    add_styled_table(doc, frontend_data[0], frontend_data[1:],
                     header_color=BRAND_TEAL,
                     col_widths=[4, 12])
    
    add_formatted_paragraph(doc, '', size=6)
    
    add_section_title(doc, '2.4 API Integration Points', 2)
    
    api_data = [
        ['API Endpoint', 'Direction', 'Purpose'],
        ['POST /api/v1/compute/launch', 'Zinda \u2192 zypl.ai', 'Request GPU instance provisioning'],
        ['POST /api/v1/compute/terminate', 'Zinda \u2192 zypl.ai', 'Terminate active GPU session'],
        ['GET /api/v1/compute/status', 'Zinda \u2192 zypl.ai', 'Check instance status and availability'],
        ['POST /api/v1/metering/report', 'zypl.ai \u2192 Zinda', 'Push real-time usage data (GPU-hours, cost)'],
        ['POST /api/v1/scoring/evaluate', 'Zinda \u2192 zypl.ai', 'Request credit score for credit model users'],
        ['GET /api/v1/pool/availability', 'Zinda \u2192 zypl.ai', 'Check GPU pool availability before showing UI'],
    ]
    
    add_styled_table(doc, api_data[0], api_data[1:],
                     header_color=BRAND_GREEN,
                     col_widths=[5, 3.5, 7.5])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════
    #  SECTION 3: DUAL PRODUCT MODELS
    # ═══════════════════════════════════════════════════════════
    
    add_section_title(doc, '3. DUAL PRODUCT MODELS', 1)
    
    add_formatted_paragraph(
        doc,
        'Clients can access GPU compute via two models: a prepaid wallet (recommended for v1 launch) '
        'and a credit limit (planned for v2 after validation). Both models operate within the same '
        '"Cloud Compute" module in the Zinda app.',
        size=10, color='#555555', space_after=10
    )
    
    add_section_title(doc, '3.1 Model A: Prepaid Wallet (\u2605 Recommended for v1)', 2)
    
    prepaid_data = [
        ['Parameter', 'Value'],
        ['How it works', 'User tops up TJS wallet \u2192 funds converted to GPU-hours at current rate \u2192 consumed on use'],
        ['Minimum top-up', 'TBD (suggested: 100 TJS / ~$10)'],
        ['Credit risk', 'Zero \u2014 fully pre-funded by the user'],
        ['Regulatory path', 'Simple \u2014 standard e-wallet/prepaid instrument'],
        ['Time to market', '2-3 months'],
        ['KYC requirements', 'Standard Zinda KYC (already completed for existing IB users)'],
    ]
    
    add_styled_table(doc, prepaid_data[0], prepaid_data[1:],
                     header_color=BRAND_GREEN,
                     col_widths=[4, 12])
    
    add_formatted_paragraph(doc, '', size=6)
    
    add_section_title(doc, '3.2 Model B: Credit Limit (Planned for v2)', 2)
    
    credit_data = [
        ['Parameter', 'Value'],
        ['How it works', 'Zinda extends a credit limit based on zypl.ai scoring \u2192 user consumes GPU-hours \u2192 pays according to fee schedule'],
        ['Credit scoring', 'zypl.ai scoring system (existing infrastructure)'],
        ['Credit limit', 'Determined per-user by scoring model'],
        ['Credit risk', 'Managed via zypl.ai scoring + usage limits + auto-terminate on limit breach'],
        ['Regulatory path', 'Requires credit product licensing from Central Bank of Tajikistan'],
        ['Time to market', '6-9 months (after v1 pilot validation)'],
        ['Max tenor', 'Up to 24 months'],
    ]
    
    add_styled_table(doc, credit_data[0], credit_data[1:],
                     header_color=BRAND_ACCENT,
                     col_widths=[4, 12])
    
    add_formatted_paragraph(doc, '', size=6)
    
    add_section_title(doc, '3.2.1 Credit Fee Schedule', 3)
    
    add_formatted_paragraph(
        doc,
        'The credit product uses a progressive fee structure. The bank fee starts at 0% for the first month '
        'and increases gradually over the tenor up to 24 months:',
        size=10, color='#555555', space_after=8
    )
    
    fee_schedule = [
        ['Month', 'Bank Fee (%)', 'Cumulative Effect', 'Notes'],
        ['Month 1', '0%', 'Free trial period', 'Zero cost to attract users and validate demand'],
        ['Month 2', '1%', 'Low entry fee', 'Minimal fee to begin revenue generation'],
        ['Month 3', '2%', 'Gradual increase', 'Encourages timely repayment'],
        ['Month 6', '3%', 'Mid-term rate', 'Standard fee for medium-term usage'],
        ['Month 9', '4%', 'Increasing incentive', 'Higher fee incentivizes earlier repayment'],
        ['Month 12', '5%', 'Annual rate', 'Maximum rate for 12-month tenor'],
        ['Month 13-24', '5% (capped)', 'Extended tenor', 'Fee remains at 5% cap through month 24'],
    ]
    
    add_styled_table(doc, fee_schedule[0], fee_schedule[1:],
                     header_color=BRAND_PURPLE,
                     col_widths=[2.5, 2.5, 4, 7])
    
    add_section_title(doc, '3.2.2 Model Comparison', 2)
    
    models_compare = [
        ['Characteristic', 'A: Prepaid Wallet\n(\u2605 RECOMMENDED for v1)', 'B: Credit Limit'],
        ['Principle', 'User deposits TJS \u2192 converts to GPU-hours \u2192 consumes',
         'Zinda extends credit line, user consumes, pays at cycle end'],
        ['Credit risk', 'None (prepaid)', 'High (requires underwriting)'],
        ['Compliance', 'Simple (standard prepaid)', 'Complex (credit product)'],
        ['User attraction', 'Medium barrier', 'Low barrier \u2014 "train now, pay later"'],
        ['Launch complexity', 'Low', 'High \u2014 needs scoring model calibration'],
        ['Speed to market', '2-3 months', '6-9 months'],
        ['Recommendation', '\u2605 RECOMMENDED for v1', 'For v2 after pilot validation'],
    ]
    
    add_styled_table(doc, models_compare[0], models_compare[1:],
                     header_color=BRAND_PRIMARY,
                     col_widths=[3.5, 6, 6.5])
    
    add_formatted_paragraph(doc, '', size=6)
    
    # Model comparison chart
    doc.add_picture(model_comp, width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════
    #  SECTION 4: CUSTOMER JOURNEY MAP (CJM)
    # ═══════════════════════════════════════════════════════════
    
    add_section_title(doc, '4. CUSTOMER JOURNEY MAP (CJM)', 1)
    
    add_formatted_paragraph(
        doc,
        'The customer journey has been fundamentally simplified by removing the physical card '
        'issuance step. All interactions now happen within the Zinda Internet Banking app. '
        'The journey covers the full cycle from discovery to retention.',
        size=10, color='#555555', space_after=12
    )
    
    add_section_title(doc, '4.1 Visual Journey Map', 2)
    doc.add_picture(cjm_chart, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_formatted_paragraph(doc, '', size=6)
    
    add_section_title(doc, '4.2 Journey Stage Details', 2)
    
    cjm_stages = [
        ['Step', 'Action', 'Actor', 'Channel', 'Detail'],
        ['1', 'Navigate to "Cloud Compute" module', 'User', 'Zinda IB App',
         'User opens the app, sees the new "Cloud Compute" section'],
        ['2', 'Accept Terms & Conditions', 'User', 'In-app modal',
         'One-time acceptance of GPU compute service terms'],
        ['3', 'Choose product model', 'User', 'In-app',
         'Select Prepaid (top-up wallet) or Credit (apply for limit)'],
        ['3a', 'Top-up wallet (Prepaid)', 'User', 'In-app transfer',
         'Transfer TJS from linked bank account to compute wallet'],
        ['3b', 'Apply for credit limit (Credit)', 'User \u2192 zypl.ai', 'Backend scoring',
         'zypl.ai scoring system evaluates and sets credit limit'],
        ['4', 'Configure & Launch Server', 'User', '"Launch Server" button',
         'Select H200 config, click "Launch Server" \u2192 triggers API call'],
        ['5', 'Backend authorization', 'System', 'Zinda \u2192 zypl.ai API',
         'Zinda checks balance/credit \u2192 calls zypl.ai provisioning endpoint'],
        ['6', 'GPU provisioned', 'zypl.ai', 'Backend',
         'H200 instance allocated, endpoint + credentials generated'],
        ['7', 'Receive access credentials', 'User', 'In-app notification',
         'SSH endpoint, API key, and Jupyter URL delivered in-app'],
        ['8', 'Use GPU', 'User', 'SSH / Jupyter / API',
         'Train models, run inference, process data on H200'],
        ['9', 'Real-time metering', 'zypl.ai \u2192 Zinda', 'API feed',
         'GPU-hours consumed, wallet balance decremented in real-time'],
        ['10', 'Session end & billing', 'System', 'Automated',
         'Session terminated (user-initiated or balance depleted), final bill settled'],
    ]
    
    add_styled_table(doc, cjm_stages[0], cjm_stages[1:],
                     header_color='#FF6F00',
                     col_widths=[1, 3.5, 2.5, 2.5, 6])
    
    add_formatted_paragraph(doc, '', size=6)
    
    add_section_title(doc, '4.3 v1.0 vs v2.0 CJM Changes', 2)
    
    cjm_changes = [
        ['Aspect', 'v1.0 (Physical Card)', 'v2.0 (Embedded Module)'],
        ['Onboarding', 'Visit branch \u2192 KYC \u2192 receive physical card', 'Open app \u2192 navigate to module \u2192 accept terms'],
        ['Authorization', 'Tap physical card at terminal', 'Click "Launch Server" button in-app'],
        ['Payment trigger', 'Card swipe/tap authorization', 'Backend API call from Zinda ledger to zypl.ai'],
        ['Delivery', 'Wait for card printing & delivery', 'Instant \u2014 already in internet banking'],
        ['User friction', 'High (branch visit, card wait, Visa compliance)', 'Very low (2-3 taps in existing app)'],
        ['Visa dependency', 'Required (BIN, 10K+ cards)', 'None'],
    ]
    
    add_styled_table(doc, cjm_changes[0], cjm_changes[1:],
                     header_color=BRAND_PURPLE,
                     col_widths=[3, 5.5, 7.5])
    
    add_formatted_paragraph(doc, '', size=6)
    
    add_section_title(doc, '4.4 Pain Points & Solutions', 2)
    
    pain_solutions = [
        ['Pain Point', 'Current Situation', 'Virtual Compute Wallet Solution'],
        ['No domestic GPU', 'No GPU-on-demand offering in Tajikistan',
         'Sovereign GPU pool (NVIDIA H200) accessible via Zinda app'],
        ['FX friction', 'Paying foreign clouds requires currency conversion + fees',
         'All transactions in TJS via Zinda wallet \u2014 zero FX'],
        ['Data sovereignty', 'Data stored on foreign servers (AWS, GCP, Azure)',
         'All infrastructure located in Tajikistan'],
        ['High cost', 'AWS/GCP/Azure: $28-35/hr for H200 equivalent',
         'Virtual Compute Wallet: $3-4/hr (85-90% savings)'],
        ['Complex access', 'Register on foreign platforms, foreign KYC, credit card',
         'One-click access in existing Zinda app \u2014 no new KYC'],
        ['Physical card barriers', 'Visa requires 10K+ cards, BIN issuance, per-card compliance',
         'No physical card needed \u2014 embedded digital module'],
    ]
    
    add_styled_table(doc, pain_solutions[0], pain_solutions[1:],
                     header_color=BRAND_ACCENT,
                     col_widths=[3, 5, 8])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════
    #  SECTION 5: BUSINESS MODEL
    # ═══════════════════════════════════════════════════════════
    
    add_section_title(doc, '5. BUSINESS MODEL', 1)
    
    add_section_title(doc, '5.1 Business Model Canvas', 2)
    
    add_formatted_paragraph(
        doc,
        'Visual representation of all business model components for the Virtual Compute Wallet '
        'following the Osterwalder methodology. Updated to reflect the embedded module architecture.',
        size=10, color='#555555', space_after=10
    )
    
    doc.add_picture(bmc_chart, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_formatted_paragraph(doc, '', size=6)
    
    # Detailed BMC table
    bmc_details = [
        ['Component', 'Description'],
        ['Key Partners',
         '\u2022 zypl.ai \u2014 GPU infrastructure provider (4\u00d7 H200), provisioning, metering, credit scoring\n'
         '\u2022 Zinda Capital \u2014 Internet banking platform, wallet operations, KYC, billing\n'
         '\u2022 NVIDIA \u2014 GPU hardware supplier\n'
         '\u2022 Central Bank of Tajikistan \u2014 Regulator'],
        ['Key Activities',
         '\u2022 GPU fleet management and instance provisioning via API\n'
         '\u2022 Real-time GPU-hour metering and wallet balance management\n'
         '\u2022 Credit scoring and limit management (v2)\n'
         '\u2022 KYC verification and regulatory compliance\n'
         '\u2022 API integration between Zinda ledger and zypl.ai infrastructure'],
        ['Value Proposition',
         '\u2022 Instant GPU access via Zinda Internet Banking app\n'
         '\u2022 TJS wallet \u2014 no FX friction, no foreign accounts\n'
         '\u2022 Pay only for GPU-hours consumed\n'
         '\u2022 Sovereign compute \u2014 data stays in Tajikistan\n'
         '\u2022 8-10x cheaper than major cloud providers'],
        ['Customer Segments',
         'MOBILE WALLET (Individuals):\n'
         '\u2022 Researchers, professors, teachers of AI\n'
         '\u2022 University students and academics\n'
         '\n'
         'INTERNET BANKING (Companies):\n'
         '\u2022 AI startups and tech companies\n'
         '\u2022 Financial companies and fintechs\n'
         '\u2022 Other enterprises'],
        ['Revenue Streams',
         '\u2022 GPU-hour billing: $3-4/hr (H200)\n'
         '\u2022 Pilot (4 GPU): $90K-$170K/year\n'
         '\u2022 Scale (100 GPU): $2.25M-$4.25M/year\n'
         '\u2022 Wallet service/maintenance fees\n'
         '\u2022 Premium SLA and priority access tiers'],
        ['Cost Structure',
         '\u2022 GPU hardware amortization ($120-160K per H200)\n'
         '\u2022 Electricity and cooling for data center\n'
         '\u2022 Zinda platform operations (IB module, wallet system)\n'
         '\u2022 zypl.ai platform development and maintenance\n'
         '\u2022 Compliance and regulatory costs'],
    ]
    
    add_styled_table(doc, bmc_details[0], bmc_details[1:],
                     header_color=BRAND_PRIMARY,
                     col_widths=[4, 12])
    
    doc.add_page_break()
    
    add_section_title(doc, '5.2 Financial Projections', 2)
    
    doc.add_picture(revenue_chart, width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_formatted_paragraph(doc, '', size=6)
    
    fin_data = [
        ['Phase', '# GPUs', 'GPU-hrs/mo', 'Min Revenue/yr', 'Max Revenue/yr', 'Notes'],
        ['Pilot', '4', '2,880', '$90K', '$170K', 'Proof of concept, validate the rail'],
        ['Phase 1', '16', '11,520', '$360K', '$680K', 'Initial scaling'],
        ['Phase 2', '50', '36,000', '$1.125M', '$2.125M', 'Regional growth'],
        ['Scale', '100', '72,000', '$2.25M', '$4.25M', 'Full deployment'],
    ]
    
    add_styled_table(doc, fin_data[0], fin_data[1:],
                     header_color=BRAND_PRIMARY,
                     col_widths=[2, 2, 2.5, 2.5, 2.8, 4.2])
    
    add_formatted_paragraph(
        doc,
        '\u203b Note: The pilot is supply-bound, not demand-bound. '
        'Revenue scales linearly with GPU count while maintaining utilization.',
        size=8, italic=True, color='#888888', space_after=10
    )
    
    add_section_title(doc, '5.3 Competitive Analysis', 2)
    
    doc.add_picture(market_chart, width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_formatted_paragraph(doc, '', size=6)
    
    comp_data = [
        ['Provider', '$/GPU-hr', 'Available in TJ', 'Data Sovereignty', 'Payment Rail'],
        ['Virtual Compute Wallet', '$3-4', '\u2705 Yes', '\u2705 Full', '\u2705 TJS wallet via Zinda app'],
        ['AWS (p4d)', '$32.77', '\u274c No', '\u274c Foreign servers', '\u26a0\ufe0f FX + intl card'],
        ['Google Cloud (A3)', '$34.72', '\u274c No', '\u274c Foreign servers', '\u26a0\ufe0f FX + intl card'],
        ['Azure (ND H200)', '$28.43', '\u274c No', '\u274c Foreign servers', '\u26a0\ufe0f FX + intl card'],
        ['Lambda Cloud', '$3.99', '\u274c No', '\u274c Foreign servers', '\u26a0\ufe0f USD only'],
    ]
    
    add_styled_table(doc, comp_data[0], comp_data[1:],
                     header_color=BRAND_ACCENT,
                     col_widths=[3, 2.5, 2.5, 3.5, 4.5])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════
    #  SECTION 6: SYSTEM FLOW DIAGRAM
    # ═══════════════════════════════════════════════════════════
    
    add_section_title(doc, '6. SYSTEM FLOW DIAGRAM', 1)
    
    add_formatted_paragraph(
        doc,
        'The diagram shows the complete system flow for the Virtual Compute Wallet \u2014 from the user\'s '
        'action in the Zinda Internet Banking app to GPU provisioning by zypl.ai and back. '
        'The flow is divided into three swimlanes: User, Zinda Capital (Ledger + API), and zypl.ai (GPU Infra).',
        size=10, color='#555555', space_after=12
    )
    
    add_section_title(doc, '6.1 Process Flow (Swimlane)', 2)
    
    doc.add_picture(flow_chart, width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_formatted_paragraph(doc, '', size=6)
    
    add_section_title(doc, '6.2 Process Step Descriptions', 2)
    
    process_steps = [
        ['Step', 'Actor', 'Action', 'Description'],
        ['1', 'User', 'Open Zinda IB App',
         'User opens the Zinda Internet Banking application on their device'],
        ['2', 'User', 'Navigate to "Cloud Compute"',
         'User navigates to the new "Cloud Compute" module/tab within the app'],
        ['3', 'User', 'Accept T&C & Fund Wallet',
         'Accepts terms (first time), then tops up TJS wallet or checks existing credit limit'],
        ['4', 'User', 'Click "Launch Server"',
         'Selects H200 GPU configuration and clicks the "Launch Server" button'],
        ['5', 'Zinda', 'Receive launch request',
         'Backend receives the "Launch Server" API request from the frontend'],
        ['6', 'Zinda', 'Check balance / credit',
         'Ledger checks if user has sufficient wallet balance or credit limit'],
        ['7', 'Zinda', 'Call zypl.ai provisioning API',
         'If funds sufficient, calls POST /api/v1/compute/launch on zypl.ai endpoint'],
        ['8', 'zypl.ai', 'Allocate GPU',
         'Allocates 1\u00d7 H200 from the 4-GPU pool, checks availability and queue'],
        ['9', 'zypl.ai', 'Provision instance',
         'Launches the H200 instance, generates SSH endpoint, API key, and Jupyter URL'],
        ['10', 'User', 'Receive credentials',
         'Gets endpoint + credentials delivered in-app. Begins working with GPU.'],
        ['11', 'zypl.ai', 'Real-time metering',
         'Tracks GPU-hours consumed in real-time, pushes metering data to Zinda via API'],
        ['12', 'Zinda', 'Debit wallet',
         'Decrements user wallet balance based on metering feed. Generates billing statements.'],
    ]
    
    add_styled_table(doc, process_steps[0], process_steps[1:],
                     header_color=BRAND_GREEN,
                     col_widths=[1, 1.8, 3.5, 9.7])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════
    #  SECTION 7: ADVANTAGES
    # ═══════════════════════════════════════════════════════════
    
    add_section_title(doc, '7. ADVANTAGES OF NEW ARCHITECTURE', 1)
    
    add_formatted_paragraph(
        doc,
        'The pivot from a physical co-branded Visa card to an embedded Internet Banking module '
        'provides significant advantages across cost, speed, compliance, and user experience:',
        size=10, color='#555555', space_after=10
    )
    
    advantages = [
        ['Advantage', 'Detail'],
        ['No Visa dependency', 'Eliminates BIN requirement, 10K card minimum, and per-card compliance certification'],
        ['Faster time-to-market', 'No card printing, no Visa certification \u2014 just a software module in existing app'],
        ['Lower cost', 'No plastic card manufacturing, no card network fees, no Visa royalties'],
        ['Better user experience', 'In-app experience is smoother than physical card workflow \u2014 2-3 taps to launch'],
        ['Existing user base', 'Leverages Zinda\'s existing internet banking users \u2014 no new KYC needed'],
        ['Flexible billing', 'Wallet-based billing is more granular than card authorization cycles'],
        ['Dual model support', 'Easy to offer both prepaid and credit within the same UI without separate card products'],
        ['Instant delivery', 'No waiting for card printing and delivery \u2014 module available immediately'],
        ['Simpler regulation', 'E-wallet/prepaid is simpler to regulate than a co-branded card product'],
    ]
    
    add_styled_table(doc, advantages[0], advantages[1:],
                     header_color=BRAND_GREEN,
                     col_widths=[4, 12])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════
    #  SECTION 8: IMPLEMENTATION ROADMAP
    # ═══════════════════════════════════════════════════════════
    
    add_section_title(doc, '8. IMPLEMENTATION ROADMAP', 1)
    
    add_section_title(doc, '8.1 Phase 1: MVP (Weeks 1-8)', 2)
    
    phase1 = [
        ['Week', 'Task', 'Owner', 'Priority'],
        ['1-2', 'Finalize product specifications (this document)', 'Zinda + zypl.ai', '\U0001f534 Critical'],
        ['2-3', 'Jointly design "Cloud Compute" module UI/UX wireframes', 'Zinda + zypl.ai', '\U0001f534 Critical'],
        ['3-4', 'Finalize GPU-hour conversion rate & revenue split agreement', 'Zinda + zypl.ai', '\U0001f534 Critical'],
        ['3-5', 'Build zypl.ai provisioning API endpoints', 'zypl.ai', '\U0001f534 Critical'],
        ['4-6', 'Integrate Zinda ledger <-> zypl.ai metering API', 'Zinda + zypl.ai', '\U0001f7e0 High'],
        ['5-7', 'Build front-end module in Zinda Internet Banking app', 'Zinda', '\U0001f7e0 High'],
        ['6-8', 'Regulatory approval for e-wallet compute product', 'Zinda', '\U0001f7e0 High'],
        ['8', 'Internal testing & QA', 'Both', '\U0001f7e1 Medium'],
    ]
    
    add_styled_table(doc, phase1[0], phase1[1:],
                     header_color=BRAND_PRIMARY,
                     col_widths=[2, 7.5, 3, 3.5])
    
    add_formatted_paragraph(doc, '', size=6)
    
    add_section_title(doc, '8.2 Phase 2: Pilot Launch (Weeks 8-12)', 2)
    
    phase2 = [
        ['Week', 'Task', 'Owner', 'Priority'],
        ['8-9', 'Soft launch with 10-20 internal/partner users', 'Both', '\U0001f7e0 High'],
        ['9-10', 'Monitor metering accuracy, billing flows, UX feedback', 'Both', '\U0001f7e0 High'],
        ['10-12', 'Iterate on UI/UX and API based on pilot feedback', 'Both', '\U0001f7e1 Medium'],
    ]
    
    add_styled_table(doc, phase2[0], phase2[1:],
                     header_color=BRAND_BLUE,
                     col_widths=[2, 7.5, 3, 3.5])
    
    add_formatted_paragraph(doc, '', size=6)
    
    add_section_title(doc, '8.3 Phase 3: Public Launch + Credit Model (Weeks 12-24)', 2)
    
    phase3 = [
        ['Week', 'Task', 'Owner', 'Priority'],
        ['12-14', 'Public launch of prepaid wallet model', 'Both', '\U0001f534 Critical'],
        ['14-18', 'Develop credit scoring integration for credit model', 'zypl.ai', '\U0001f7e0 High'],
        ['18-20', 'Regulatory approval for credit product', 'Zinda', '\U0001f7e0 High'],
        ['20-24', 'Launch credit limit model (v2)', 'Both', '\U0001f7e1 Medium'],
    ]
    
    add_styled_table(doc, phase3[0], phase3[1:],
                     header_color=BRAND_PURPLE,
                     col_widths=[2, 7.5, 3, 3.5])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════
    #  SECTION 9: OPEN QUESTIONS & NEXT STEPS
    # ═══════════════════════════════════════════════════════════
    
    add_section_title(doc, '9. OPEN QUESTIONS & NEXT STEPS', 1)
    
    add_section_title(doc, '9.1 Open Questions', 2)
    
    questions = [
        ['#', 'Question', 'Impact', 'Status'],
        ['1', 'What is the agreed revenue split between Zinda (wallet/billing) and zypl.ai (infrastructure)?',
         'Critical \u2014 determines financial viability for both parties', 'To be agreed'],
        ['2', 'What is the exact TJS \u2192 GPU-hour conversion rate? Fixed or floating with USD exchange rate?',
         'Critical \u2014 affects user pricing and billing accuracy', 'To be decided'],
        ['3', 'How is the "Cloud Compute" wallet classified under Tajik law? E-money? Prepaid instrument? Service-linked wallet?',
         'High \u2014 determines regulatory path and licensing requirements', 'To be clarified with regulator'],
        ['4', 'For the credit model (v2), will zypl.ai use existing consumer scoring or develop a GPU-specific scoring model?',
         'Medium \u2014 affects credit model accuracy and risk', 'To be decided'],
        ['5', 'What are the minimum and maximum wallet top-up amounts?',
         'Medium \u2014 affects user accessibility and risk management', 'To be decided'],
    ]
    
    add_styled_table(doc, questions[0], questions[1:],
                     header_color=BRAND_ACCENT,
                     col_widths=[1, 6.5, 4.5, 4])
    
    add_formatted_paragraph(doc, '', size=8)
    
    add_section_title(doc, '9.2 Immediate Next Steps', 2)
    
    next_steps = [
        ['#', 'Action', 'Responsible', 'Deadline'],
        ['1', 'Finalize and sign off on product specifications (this document)', 'Zinda + zypl.ai', 'Week 1'],
        ['2', 'Jointly discuss and design product UI/UX (wireframes, user flows, branding)', 'Zinda + zypl.ai', 'Week 2-3'],
        ['3', 'Agree on GPU-hour conversion rate and revenue split model', 'Zinda + zypl.ai', 'Week 2-3'],
        ['4', 'Confirm regulatory classification of the compute wallet with Central Bank', 'Zinda Capital', 'Week 3-4'],
        ['5', 'Build and document zypl.ai provisioning API endpoints', 'zypl.ai', 'Week 3-5'],
        ['6', 'Establish API integration testing environment', 'Both', 'Week 4-5'],
        ['7', 'Begin front-end development of the embedded module', 'Zinda', 'Week 5-7'],
    ]
    
    add_styled_table(doc, next_steps[0], next_steps[1:],
                     header_color=BRAND_PRIMARY,
                     col_widths=[1, 7, 3.5, 4.5])
    
    add_formatted_paragraph(doc, '', size=10)
    
    # ═══════════════════════════════════════════════════════════
    #  CLOSING
    # ═══════════════════════════════════════════════════════════
    
    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2501' * 40)
    run.font.color.rgb = RGBColor(*hex_to_rgb(BRAND_ACCENT))
    run.font.size = Pt(12)
    
    add_formatted_paragraph(
        doc, 'Document prepared: June 2, 2026 | v2.0 (Post-meeting update)', size=9, italic=True,
        color='#999999', alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=10
    )
    add_formatted_paragraph(
        doc, 'Virtual Compute Wallet \u2014 zypl.ai \u00d7 Zinda Capital', size=10, bold=True,
        color=BRAND_PRIMARY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )
    add_formatted_paragraph(
        doc, 'Embedded GPU Credit Module | Zinda Internet Banking', size=9, italic=False,
        color=BRAND_PURPLE, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )
    add_formatted_paragraph(
        doc, 'Confidential | For internal use only', size=8,
        italic=True, color='#BBBBBB', alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    
    # Save document
    output_path = os.path.join(OUTPUT_DIR, 'Virtual_Compute_Wallet_Product_Specs_v2.docx')
    doc.save(output_path)
    print(f"\n[DONE] Document saved: {output_path}")
    print(f"[DONE] Charts saved in: {CHARTS_DIR}")
    
    # Also save a copy with the old name for backwards compatibility
    compat_path = os.path.join(OUTPUT_DIR, 'GPU_Credit_Card_Business_Document.docx')
    doc.save(compat_path)
    print(f"[DONE] Compatibility copy saved: {compat_path}")
    
    return output_path


if __name__ == '__main__':
    generate_document()
