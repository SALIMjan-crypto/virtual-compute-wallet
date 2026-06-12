# -*- coding: utf-8 -*-
"""
Virtual Compute Wallet — 1-Page Product Specifications Summary
Generates a single-page Word document covering all key product aspects.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = r"c:\ANTIGRAVITY PROJECTS\GPU Credit Card"

# Colors
NAVY = "#1A1A2E"
RED = "#E94560"
GREEN = "#00B894"
BLUE = "#0984E3"
PURPLE = "#6C5CE7"
GOLD = "#D4A017"
DARK_GRAY = "#333333"
MED_GRAY = "#555555"
LIGHT_GRAY = "#888888"


def hex_to_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def set_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex.lstrip("#")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_run(paragraph, text, size=8, bold=False, italic=False, color=None, font='Calibri'):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        r, g, b = hex_to_rgb(color)
        run.font.color.rgb = RGBColor(r, g, b)
    return run


def make_paragraph(doc, space_after=2, space_before=0, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = Pt(11)
    return p


def add_compact_table(doc, headers, rows, header_color=NAVY, col_widths=None, font_size=7):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Reduce table cell margins
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            margins = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'<w:top w:w="20" w:type="dxa"/>'
                f'<w:left w:w="40" w:type="dxa"/>'
                f'<w:bottom w:w="20" w:type="dxa"/>'
                f'<w:right w:w="40" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tcPr.append(margins)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(h)
        run.font.name = 'Calibri'
        run.font.size = Pt(font_size - 0.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_shading(cell, header_color)

    # Rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[1 + ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(*hex_to_rgb(DARK_GRAY))
            if ri % 2 == 0:
                set_shading(cell, "#F0F3FF")

    return table


def section_header(doc, text, color=NAVY):
    p = make_paragraph(doc, space_before=6, space_after=2)
    add_run(p, '\u2588\u2588 ', size=8, color=color, bold=True)
    add_run(p, text, size=9, bold=True, color=color)
    return p


def generate_one_pager():
    doc = Document()

    # Page setup — A4, tight margins for 1-page fit
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(8)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = Pt(11)

    # ════════════════════════════════════════════════════════
    # TITLE
    # ════════════════════════════════════════════════════════
    p = make_paragraph(doc, space_after=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, 'VIRTUAL COMPUTE WALLET', size=18, bold=True, color=NAVY)

    p = make_paragraph(doc, space_after=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, 'Embedded GPU Credit Module  ', size=9, bold=True, color=PURPLE)
    add_run(p, '|  ', size=9, color=LIGHT_GRAY)
    add_run(p, 'zypl.ai \u00d7 Zinda Capital', size=9, bold=True, color=RED)

    p = make_paragraph(doc, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, 'Product Specifications v2.0  |  June 2, 2026  |  Post-meeting update (Akai Akmal / Visa)',
            size=7, italic=True, color=LIGHT_GRAY)

    # Divider
    p = make_paragraph(doc, space_after=3, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, '\u2501' * 80, size=6, color=RED)

    # ════════════════════════════════════════════════════════
    # PIVOT NOTICE
    # ════════════════════════════════════════════════════════
    p = make_paragraph(doc, space_after=4)
    add_run(p, '\u26a1 PRODUCT PIVOT: ', size=7.5, bold=True, color=RED)
    add_run(p, 'Visa requires 10,000+ cards for a dedicated BIN + per-card compliance \u2014 impractical for pilot. ', size=7, color=MED_GRAY)
    add_run(p, 'New architecture: ', size=7, bold=True, color=DARK_GRAY)
    add_run(p, 'Embedded module inside Zinda Internet Banking as a ', size=7, color=MED_GRAY)
    add_run(p, 'purpose-locked TJS wallet', size=7, bold=True, color=RED)
    add_run(p, '. Funds can ONLY be used for GPU compute. No physical card, no Visa dependency.', size=7, color=MED_GRAY)

    # ════════════════════════════════════════════════════════
    # ARCHITECTURE & TECH SPECS
    # ════════════════════════════════════════════════════════
    section_header(doc, 'ARCHITECTURE & TECHNICAL SPECIFICATIONS', BLUE)

    add_compact_table(doc,
        ['Component', 'Specification'],
        [
            ['Product Type', 'Embedded "Cloud Compute" module in Zinda Internet Banking app'],
            ['Backend', 'Dedicated TJS e-wallet on Zinda\'s ledger (purpose-locked for GPU compute only)'],
            ['API Integration', 'Zinda ledger \u2194 zypl.ai provisioning endpoint (REST API)'],
            ['GPU Hardware', 'NVIDIA H200 SXM \u2014 141 GB HBM3e per GPU \u2014 Hopper architecture'],
            ['Pilot Pool', '4\u00d7 H200 (1 server) \u2014 2,880 GPU-hours/month \u2014 4 concurrent sessions max'],
            ['Pricing', '$3\u20134 / GPU-hour (vs $28\u201335 on AWS/GCP/Azure \u2014 85\u201390% savings)'],
            ['Billing', 'Real-time metering via zypl.ai API \u2192 automatic wallet debit'],
        ],
        header_color=BLUE, col_widths=[3.5, 14.5])

    # ════════════════════════════════════════════════════════
    # CUSTOMER SEGMENTS
    # ════════════════════════════════════════════════════════
    section_header(doc, 'CUSTOMER SEGMENTS BY CHANNEL', PURPLE)

    add_compact_table(doc,
        ['Channel', 'Target Users', 'Access Method'],
        [
            ['Mobile Wallet App', 'Individual researchers, professors, university teachers of AI, students', 'Dedicated mobile wallet app with purpose-locked GPU funds'],
            ['Internet Banking', 'AI startups, tech companies, financial companies, fintechs, enterprises', 'Dedicated e-wallet within Zinda Internet Banking platform'],
        ],
        header_color=PURPLE, col_widths=[3.5, 8, 6.5])

    # ════════════════════════════════════════════════════════
    # DUAL PRODUCT MODELS
    # ════════════════════════════════════════════════════════
    section_header(doc, 'DUAL PRODUCT MODELS', GREEN)

    add_compact_table(doc,
        ['', 'Model A: Prepaid Wallet  \u2605 RECOMMENDED v1', 'Model B: Credit Limit  (Planned v2)'],
        [
            ['How it works', 'User tops up TJS wallet \u2192 converts to GPU-hours \u2192 consumed on use', 'Zinda extends credit (scored by zypl.ai) \u2192 user consumes \u2192 pays per fee schedule'],
            ['Credit risk', 'Zero \u2014 fully pre-funded', 'Managed by zypl.ai scoring'],
            ['Bank fee', '0% \u2014 no interest, pure prepaid', 'Progressive: 0% (M1) \u2192 1% (M2) \u2192 2% (M3) \u2192 3% (M6) \u2192 4% (M9) \u2192 5% (M12\u201324)'],
            ['Max tenor', 'N/A', 'Up to 24 months'],
            ['Max tenor', 'N/A', 'Up to 24 months'],
            ['Regulatory', 'Simple \u2014 standard e-wallet', 'Complex \u2014 requires credit product license'],
            ['Time to market', '2\u20133 months', '6\u20139 months (after v1 validation)'],
        ],
        header_color=GREEN, col_widths=[3, 7.5, 7.5])

    # ════════════════════════════════════════════════════════
    # DEFAULT GUARANTEE & FEE SCHEDULE
    # ════════════════════════════════════════════════════════
    section_header(doc, 'CREDIT FEE SCHEDULE', RED)

    p = make_paragraph(doc, space_after=2)
    add_run(p, 'zypl.ai\'s scoring model pre-screens users to minimize defaults; GPU sessions auto-terminate on limit breach.', size=7, color=MED_GRAY)

    add_compact_table(doc,
        ['Month 1', 'Month 2', 'Month 3', 'Month 6', 'Month 9', 'Month 12', 'Month 13\u201324'],
        [
            ['0% Free', '1%', '2%', '3%', '4%', '5%', '5% (capped)'],
        ],
        header_color=PURPLE, col_widths=[2.57, 2.57, 2.57, 2.57, 2.57, 2.57, 2.57])

    # ════════════════════════════════════════════════════════
    # CUSTOMER JOURNEY
    # ════════════════════════════════════════════════════════
    section_header(doc, 'CUSTOMER JOURNEY (IN-APP)', BLUE)

    p = make_paragraph(doc, space_after=3)
    steps = [
        ('1', 'Open Zinda App \u2192 "Cloud Compute"'),
        ('2', 'Accept T&C & Fund Wallet (TJS)'),
        ('3', 'Configure H200 GPU instance'),
        ('4', 'Click "Launch Server"'),
        ('5', 'Receive SSH + Jupyter credentials'),
        ('6', 'Train models \u2014 real-time billing'),
    ]
    for i, (num, text) in enumerate(steps):
        add_run(p, f' {num} ', size=7.5, bold=True, color='#FFFFFF')
        # We can't easily do a circle background in docx, so use brackets
        add_run(p, f' {text}', size=7, color=DARK_GRAY)
        if i < len(steps) - 1:
            add_run(p, '  \u2192  ', size=8, color=LIGHT_GRAY)

    p = make_paragraph(doc, space_after=2)
    add_run(p, 'No branch visit  \u2022  No physical card  \u2022  No Visa  \u2022  2\u20133 taps to launch  \u2022  Purpose-locked wallet',
            size=6.5, italic=True, color=LIGHT_GRAY)

    # ════════════════════════════════════════════════════════
    # REVENUE PROJECTION
    # ════════════════════════════════════════════════════════
    section_header(doc, 'REVENUE PROJECTION & PARTICIPANTS', NAVY)

    add_compact_table(doc,
        ['Phase', 'GPUs', 'GPU-hrs/mo', 'Revenue/yr', 'Zinda Capital', 'zypl.ai'],
        [
            ['Pilot', '4', '2,880', '$90K\u2013$170K', 'Wallet issuer, KYC, billing, compliance', 'GPU infra (4\u00d7 H200), provisioning API, metering'],
            ['Phase 1', '16', '11,520', '$360K\u2013$680K', 'Credit underwriting, regulatory', 'Credit scoring'],
            ['Phase 2', '50', '36,000', '$1.1M\u2013$2.1M', '', ''],
            ['Scale', '100', '72,000', '$2.25M\u2013$4.25M', '', ''],
        ],
        header_color=NAVY, col_widths=[2, 1.5, 2.5, 3, 4.5, 4.5])

    # ════════════════════════════════════════════════════════
    # NEXT STEPS
    # ════════════════════════════════════════════════════════
    section_header(doc, 'IMMEDIATE NEXT STEPS', GREEN)

    add_compact_table(doc,
        ['#', 'Action', 'Responsible', 'Timeline'],
        [
            ['1', 'Finalize & sign off on product specifications (this document)', 'Zinda + zypl.ai', 'Week 1'],
            ['2', 'Jointly discuss & design product UI/UX (wireframes, user flows, branding)', 'Zinda + zypl.ai', 'Week 2\u20133'],
            ['3', 'Agree on GPU-hour conversion rate & revenue split', 'Zinda + zypl.ai', 'Week 2\u20133'],
            ['4', 'Build zypl.ai provisioning API + Zinda ledger integration', 'zypl.ai + Zinda', 'Week 3\u20135'],
            ['5', 'Regulatory filing for e-wallet compute product', 'Zinda Capital', 'Week 3\u20134'],
            ['6', 'Front-end development, QA & pilot launch (10\u201320 users)', 'Both', 'Week 5\u20138'],
        ],
        header_color=GREEN, col_widths=[1, 9, 3.5, 4.5])

    # ════════════════════════════════════════════════════════
    # FOOTER
    # ════════════════════════════════════════════════════════
    p = make_paragraph(doc, space_before=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, '\u2501' * 80, size=5, color=RED)

    p = make_paragraph(doc, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, 'CONFIDENTIAL', size=6, bold=True, color=RED)
    add_run(p, '  |  Virtual Compute Wallet  |  zypl.ai \u00d7 Zinda Capital  |  v2.0  |  June 2026', size=6, color=LIGHT_GRAY)

    # Save
    path = os.path.join(OUTPUT_DIR, 'Virtual_Compute_Wallet_One_Pager.docx')
    doc.save(path)
    print(f"[DONE] One-pager saved: {path}")
    return path


if __name__ == '__main__':
    generate_one_pager()
